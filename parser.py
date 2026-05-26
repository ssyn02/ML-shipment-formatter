from flask import json
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta, UTC
from collections import defaultdict
import pandas as pd
import requests
import sys


def resolve_product_id(mla, token, item_cache):

    if mla not in item_cache:
        response = requests.get(
            f"https://api.mercadolibre.com/items/{mla}",
            headers={
                "Authorization": f"Bearer {token}"
            }
        ).json()

        if not str(response["user_product_id"]):
            item_cache[mla] = str(response["variations"]["id"])
        else:
            item_cache[mla] = str(response["user_product_id"])[3:]

    return item_cache[mla]


def fetch_shipment(shipment_id, access_token):
    return requests.get(
        f"https://api.mercadolibre.com/shipments/{shipment_id}",
        headers={"Authorization": f"Bearer {access_token}"}
    ).json()


def validate(shipment):
    return (shipment["logistic_type"] != "fulfillment" and 
            shipment["substatus"] == "ready_to_print")


def run_parser(tokens):
    user_id = 123456789  # reemplazar con tu user id
    offset = 0
    limit = 50
    date_from = (datetime.now(UTC) - timedelta(days=4)).strftime("%Y-%m-%dT%H:%M:%S.000-00:00")

    LOGISTIC_TYPES = {
        "self_service": "Flex",
        "cross_docking": "Colecta",
    }

    print("Inicializando procesamiento de ventas, no cierre esta ventana hasta terminar.\nLa informacion de salida se encontrara en la carpeta \"output\".")

    og_stdout = sys.stdout
    og_stderr = sys.stderr
    sys.stdout = open("output/log.txt", "w", encoding="utf-8")
    sys.stderr = sys.stdout

    df = pd.read_excel("catalog/productos.xlsx", sheet_name="Publicaciones")
    duplicated = set()
    grouped = defaultdict(list)
    session = requests.Session()
    item_cache = orders_info = catalog = {}
    current_key = None

    for _, row in df.iterrows():

        mla = row["mla"]
        product_id = row["product_id"]
        description = row["descripcion"]

        if pd.notna(mla):

            current_key = (
                str(mla).strip(),
                str(product_id).strip()
            )

            catalog[current_key] = [description]
        else:

            if current_key is not None:
                catalog[current_key].append(
                    description
                )

    while True:
        orders = session.get(
            f"https://api.mercadolibre.com/orders/search",
            params={
                "seller": user_id,
                "sort": "date_desc",
                "limit": limit,
                "offset": offset,
                "order.date_created.from": date_from
            },
            headers={
                "Authorization": f"Bearer {tokens['access_token']}",
                "x-format-new": "true"
            }
        ).json()["results"]

        if not orders:
            break

        shipment_ids = []
        for order in orders:
            shipping = order.get("shipping")

            if not shipping:
                continue

            shipment_id = shipping["id"]

            if shipment_id not in orders_info:
                orders_info[shipment_id] = []

            orders_info[shipment_id].extend(order.get("order_items", []))

            if shipment_id in duplicated:
                continue

            duplicated.add(shipment_id)
            shipment_ids.append(shipment_id)

        with ThreadPoolExecutor(max_workers=10) as executor:
            results = executor.map(fetch_shipment, shipment_ids, 
                                [tokens['access_token']] * len(shipment_ids))
            
            for shipment in results:
                if validate(shipment):
                    shipment["internal_items"] = []
                    rec = shipment["receiver_address"]["receiver_name"]
                    order_items = orders_info.get(shipment["id"], [])
                    print(shipment["id"])           # log para debugging, se puede sacar
                    print(shipment.get("order_id"))
                    print(shipment.get("pack_id"))
                    print(rec)

                    for item in order_items:
                        mla = item["item"]["id"]
                        product_id = item["item"].get("variation_id")
                        quantity = item["quantity"]

                        if not product_id:
                            product_id = resolve_product_id(mla, tokens['access_token'], item_cache)

                        descriptions = catalog.get((mla, str(product_id)), [f"NO ENCONTRADO: {mla}"])
                        for item in descriptions:
                            if item.startswith("NO ENCONTRADO"):
                                shipment["internal_items"].append(item)
                                print((mla, product_id), ":", item, rec)
                                continue
                            parts = item.split(" ", 1)
                            final_quantity = int(parts[0]) * quantity
                            shipment["internal_items"].append(f"{final_quantity} {parts[1]}")
                            print((mla, product_id), ":", f"{final_quantity} {parts[1]}", rec)

                    grouped[shipment["logistic_type"]].append(shipment)

        print("Procesados: ", len(grouped['self_service']) + len(grouped['cross_docking']))
        offset += limit

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    for logistic_type, shipments in grouped.items():
        document.add_paragraph(LOGISTIC_TYPES[logistic_type].capitalize())

        for shipment in shipments:
            for item in shipment.get("internal_items", []):
                prods = document.add_paragraph()
                prods.add_run(item).bold = True
            recipient = document.add_paragraph()
            recipient_name = (
                shipment.get("receiver_address", {})
                .get("receiver_name")
                or "SIN NOMBRE"
            )
            recipient.add_run(recipient_name + "\n\n")

    document.save("output/shipments.docx")
    print(f"Documento creado en carpeta output.")
    sys.stdout = og_stdout
    sys.stderr = og_stderr
